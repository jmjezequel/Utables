import docx # requires pip install python-docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from numbers import Number
from ioformats.filerw import FileWriter
from ioformats.writers import SkipWriter
from ioformats import availableWriters,TEXT,TABLE,BIBLIOGRAPHY

class DocxWriter(FileWriter):
    def __init__(self,numbered=False,outputDir='.',multiSheetOutput=False,editMode=False):
        super().__init__(numbered,outputDir,multiSheetOutput,editMode)
        self.extension = '.docx'
        self.subwriter = TextSubwriter(self)
        
    def _getopendoc(self,name=None):
        return docx.Document(name)
    
    def _savedoc(self,filename):
        self.doc.save(filename)

    def openSheet(self,sheetname,sheetType=TEXT,**kwargs):
        super().openSheet(sheetname,sheetType,**kwargs)
        self.subwriter = subwriters.get(sheetType,SkipWriter)(self)
        self.subwriter.openSheet(sheetname,sheetType,**kwargs)

    def closeSheet(self):
        if self.editMode:
            self.subwriter.closeSheet()
        super().closeSheet()

    def _deleteBetween(self,startPar,endPar): #TODO find another method because this one is is grossly inefficient
        tobedeleted = self._getElementsBetween(startPar.text,endPar.text)
        if tobedeleted != None:
            for par in tobedeleted:
                _deleteElement(par)

    def _getElementsBetween(self,startTag,endtag):
        '''return list of elements between the two marks, non including them, ie ]startTag,endtag['''
        result = None
        for p in self.doc.paragraphs:
            text = p.text
            if result != None: # we have already found start tag
                if text==endtag:
                    return result
                result.append(p)
            elif text==startTag:
                result = []
        return None # endtag not found, return None


    def _lookForParagraph(self,key):
        ''' look for a paragraph starting with key'''
        for p in self.doc.paragraphs:
            if p.text.startswith(key):
                return p
        return None
                
    def _lookForTable(self,key):
        ''' look for a table whose cell[0,0] starts with key'''
        for t in self.doc.tables:
            c = t.cell(0,0)
            if c != None and c.text.startswith(key):
                return t
        return None
                
    def writeTitle(self,element,** kwargs): # always=false,level=1,insertMode=False,style=None
        self.subwriter.writeTitle(element,** kwargs)

    def startNewLine(self): 
        self.subwriter.startNewLine()
    
    def append(self,element,** kwargs):
        self.subwriter.append(element,** kwargs)

    def writeln(self,iterable,** kwargs):
        self.subwriter.writeln(iterable, ** kwargs)

def _makeParagraphInvisible(par):
    for run in par.runs:
        run.font.hidden = True

def _deleteElement(elem):
        p = elem._element
        p.getparent().remove(p)
        p._p = p._element = None

class TableSubwriter():
    def __init__(self,parent=None):
        self.parent = parent
        self.table = None
        self.kwargs = None
        self.oldtable = None
        self.width = 0
        self.startMark = None
        self.currentrow = None
        self.col = 0
        
    def openSheet(self,sheetname,sheetType,**kwargs):
        self.kwargs = kwargs
        if self.parent.editMode: # sheetname is interpreted as the placeholder in the doc where new text should be inserted
            self.oldtable = self.parent._lookForTable(sheetname) 
            if self.oldtable == None:
                self.parent.subwriter = SkipWriter() # TODO: log some error
                return
        else: # prepare for further edit in the future
            self.parent.doc.add_heading(sheetname, 3) # add at level 3

    def closeSheet(self):
        if self.oldtable != None:
            _deleteElement(self.oldtable)
            
    def _setupTable(self,iterable):
        if self.table == None:
            row = list(iterable)
            self.width = len(row)
            self.table = self.parent.doc.add_table(0,self.width)
#            self.table.style = 'LightShadingAccent1'
            if self.parent.editMode: #mv table to startMark
                self.oldtable._tbl.addnext(self.table._tbl) 
                self.table.alignment = self.oldtable.alignment
                self.table.style = self.oldtable.style
                self.table.autofit = self.oldtable.autofit
            else:
                self.table.autofit = True
                for style,value in self.kwargs.items():
                    setattr(self.table, style, value)
            return row
        return iterable
    
    def getLine(self):
        return self.parent.getCurrentLine()
    
    def startNewLine(self):
        self.currentrow = self.table.add_row()
        self.parent._incLineCount()
        self.col = 0
        
    def _writeNewRow(self,iterable,** kwargs):
        self.startNewLine()
        for e in iterable:
            self._writeCell(e,** kwargs)
            self.col += 1
            
    def _writeCell(self,element,** kwargs):
        cell = self.currentrow.cells[self.col]
        par = cell.paragraphs[0]
        if self.getLine() == 0 and self.col == 0: # Insert Mark
            mark = par.add_run(self.parent.sheetName)
            mark.font.hidden = True
        if self.oldtable != None:
            par.style = self.oldtable.cell(0,0).paragraphs[0].style
            try:
                oldcell = self.oldtable.cell(self.getLine(),self.col)
            except:  # try first col
                if self.col == 0:
                    oldcell = self.oldtable.cell(0,0)
                else:
                    try:
                        oldcell = self.oldtable.cell(0,1)
                    except:
                        oldcell = self.oldtable.cell(0,0)                    
            cell.width = oldcell.width    
        # par.paragraph_format.line_spacing = 0.75
        if self.getLine() > 0 and isinstance(element, Number):
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = par.add_run(str(element))
        for style,value in kwargs.items():
            setattr(run, style, value)
    
    def writeTitle(self,element,** kwargs): 
        ''' write a title line onto table'''
        iter = self._setupTable(element)
        if not self.parent.editMode:
            kwargs['bold'] = True
        self._writeNewRow(iter,** kwargs)
        
    def append(self,element,** kwargs):
        if self.col >= self.width:
            self.startNewLine()
        self._writeCell(element, **kwargs)
        self.col += 1
    
    def writeln(self,iterable,** kwargs):
        iter = self._setupTable(iterable)
        self._writeNewRow(iter,** kwargs)

class TextSubwriter():
    def __init__(self,parent=None):
        self.parent = parent
        self.paragraph = None
        self.kwargs = {}
        
    def openSheet(self,sheetname,sheetType,** kwargs):
        self.kwargs = kwargs
        if self.parent.editMode: # sheetname is interpreted as the placeholder in the doc where new text should be inserted
            self.startMark = self.parent._lookForParagraph(sheetname) 
            if self.startMark == None:
                self.parent.subwriter = SkipWriter() # TODO: log some error
            else: # prepare for further edit in the future
                self.endMark = self.parent._lookForParagraph(sheetname+'.END')
                if self.endMark == None: #First insertion, add an end point marker
                    self.endMark = self.startMark
                    run = self.endMark.add_run('.END')
                    _makeParagraphInvisible(self.endMark)
                    self.startMark =  self.endMark.insert_paragraph_before(sheetname,style=self.endMark.style) # put back start marker
                    _makeParagraphInvisible(self.startMark)
                else: #must first remove paragraphs between startMark and endMark
                    self.parent._deleteBetween(self.startMark,self.endMark)
        else:
            # add a page break to start a new page 
            # self.doc.add_page_break()
            self.parent.doc.add_heading(sheetname, 3) # add at level 1

    def closeSheet(self):
        self.startMark.paragraph_format.keep_with_next = True #to prevent unfortunate separation on page breaks
    
    def writeTitle(self,element,** kwargs):
        """ write a title line onto paragraph"""
        if self.parent.editMode:
            #caveat can only add_heading at the end of doc
            self.startNewLine()
            self.append(element,** kwargs)
        else:
            self.parent.doc.add_heading(element,kwargs.get('level',1))

    def append(self,element,** kwargs):
        self._addRunLike(element,**kwargs)

    def writeln(self,iterable,** kwargs):
        self.startNewLine()
        if self.parent.isNumbered():
            self._addRunLike(self.parent.getLinePrefix(),**kwargs)
        first = True
        for elem in iterable:
            e = elem if first else ' '+elem
            first = False
            self._addRunLike(e, **kwargs)
            
    def startNewLine(self):
        if self.parent.editMode:
            self.paragraph = self.endMark.insert_paragraph_before()
            self.paragraph.style = self.endMark.style
            self.runmodel = self.endMark.runs[0]
        else: # add at the end of doc
            self.paragraph = self.parent.doc.add_paragraph()
            self.runmodel = None
        self.parent._incLineCount()


    def _addRunLike(self, elem, **kwargs):
        run = self.paragraph.add_run(elem)
        runmodel = self.runmodel
        if runmodel is not None:
            run.bold = runmodel.bold
            run.italic = runmodel.italic
            run.underline = runmodel.underline
            run.font.color.rgb = runmodel.font.color.rgb
            run.style.name = runmodel.style.name
        for style,value in kwargs.items():
            setattr(run, style, value)
        return run

class BiblioSubwriter(TextSubwriter):
    def __init__(self,parent=None):
        super().__init__(parent)
        
    def writeTitle(self,element,** kwargs): 
        ''' write a title line onto paragraph'''
        l = self.parent.getCurrentLine()
        super().writeTitle(element,** kwargs)
        self.parent.setLineNumber(l-1) #don't count titles in the numbering 

    def writeln(self,publication,** kwargs):
        self.startNewLine()
        if self.parent.isNumbered():
            self._addRunLike(self.parent.getLinePrefix(),**kwargs)
        publication.write(self,**self.kwargs) # publication is going to call back append below

    def append(self,element, venue=False, ** kwargs):
        if venue:
            self._addRunLike("In ")
            self._addRunLike(element,italic=True)
        else:
            self._addRunLike(element,**kwargs)


availableWriters['docx'] = DocxWriter()
subwriters = {TEXT:TextSubwriter,TABLE:TableSubwriter,BIBLIOGRAPHY:BiblioSubwriter}
